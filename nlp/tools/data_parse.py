import docx
from docx import Document

def split_document_into_chunks(file_path, chunk_size=100):
    doc = Document(file_path)
    text = ''.join([para.text for para in doc.paragraphs])
    
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    return chunks

def create_new_word_document(text_chunks, output_file):
    doc = docx.Document()
    for chunk in text_chunks:
        doc.add_paragraph(chunk)
    doc.save(output_file)
